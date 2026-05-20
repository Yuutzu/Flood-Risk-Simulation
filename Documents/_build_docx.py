"""Generate User_Guide.docx from USER_GUIDE.md (lightweight markdown subset).

Run:
  python "Documents/_build_docx.py"
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

BASE = Path(__file__).resolve().parent.parent
MD   = BASE / "Documents" / "USER_GUIDE.md"
OUT  = BASE / "Documents" / "User_Guide.docx"

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY  = RGBColor(0x55, 0x55, 0x55)


def set_cell_border(cell, color="000000", sz=4):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(sz))
        b.set(qn("w:color"), color)
        borders.append(b)
    tc_pr.append(borders)


def shade_cell(cell, fill="000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = BLACK


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = BLACK


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = BLACK


def add_para(doc, text, *, mono=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    if mono:
        r.font.name = "Consolas"
    if italic:
        r.italic = True


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(11)


def add_code(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line if line else " ")
        r.font.name = "Consolas"
        r.font.size = Pt(10)


def add_table(doc, headers, rows):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.autofit = True
    for c, h in enumerate(headers):
        cell = table.rows[0].cells[c]
        cell.text = ""
        shade_cell(cell, "000000")
        set_cell_border(cell)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            set_cell_border(cell, color="888888")
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(9)


def parse_md_table(lines, i):
    """Parse a github-style markdown table starting at line i. Returns (headers, rows, next_i)."""
    header_line = lines[i]
    sep_line    = lines[i + 1]
    if "|" not in header_line or not re.match(r"\s*\|?\s*[-:|\s]+\|?\s*$", sep_line):
        return None
    def split(s):
        s = s.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        return [c.strip() for c in s.split("|")]
    headers = split(header_line)
    j = i + 2
    rows = []
    while j < len(lines) and "|" in lines[j] and lines[j].strip():
        rows.append(split(lines[j]))
        j += 1
    return headers, rows, j


def build_docx_from_md(md_path: Path, out_path: Path) -> None:
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.left_margin   = Inches(0.9)
        section.right_margin  = Inches(0.9)
        section.top_margin    = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("JADE VALLEY FLOOD SIMULATION")
    r.bold = True
    r.font.size = Pt(28)
    sub = doc.add_paragraph()
    rs = sub.add_run("User Guide  |  Davao City, Philippines")
    rs.font.size = Pt(13)
    rs.italic = True
    doc.add_paragraph()  # spacer

    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    in_code = False
    code_buf: list = []
    i = 0
    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.strip().startswith("```"):
            if in_code:
                add_code(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # Skip the very first H1 - we already wrote the title block
        if line.startswith("# ") and i < 5:
            i += 1
            continue

        # Tables
        if "|" in line and i + 1 < len(lines) \
           and re.match(r"\s*\|?\s*[-:|\s]+\|?\s*$", lines[i + 1]):
            res = parse_md_table(lines, i)
            if res is not None:
                headers, rows, next_i = res
                add_table(doc, headers, rows)
                doc.add_paragraph()
                i = next_i
                continue

        # Headings
        if line.startswith("### "):
            add_h3(doc, line[4:].strip())
        elif line.startswith("## "):
            add_h2(doc, line[3:].strip())
        elif line.startswith("# "):
            add_h1(doc, line[2:].strip())
        elif re.match(r"^---+\s*$", line):
            # Horizontal rule
            p = doc.add_paragraph()
            pr = p.paragraph_format
            pr.space_before = Pt(2)
            pr.space_after = Pt(2)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "000000")
            pBdr.append(bot)
            pPr.append(pBdr)
        elif line.strip().startswith("- "):
            add_bullet(doc, line.strip()[2:])
        elif line.strip().startswith("* "):
            add_bullet(doc, line.strip()[2:])
        elif line.strip().startswith("*") and line.strip().endswith("*") \
                and len(line.strip()) > 2:
            add_para(doc, line.strip().strip("*"), italic=True)
        elif line.strip() == "":
            # blank line as paragraph break (skip extras)
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_paragraph()
        else:
            # Strip simple inline backticks
            txt = re.sub(r"`([^`]+)`", r"\1", line)
            add_para(doc, txt)
        i += 1

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    if not MD.exists():
        raise SystemExit(f"Source markdown not found: {MD}")
    build_docx_from_md(MD, OUT)
